"""Helpers for generating document fixtures in tests."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfWriter

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures" / "documents"


def sample_txt_bytes() -> bytes:
    return (FIXTURES_DIR / "sample.txt").read_bytes()


def sample_md_bytes() -> bytes:
    return (FIXTURES_DIR / "sample.md").read_bytes()


def make_pdf_bytes(text: str = "Cortexa is a local-first AI agent platform.") -> bytes:
    """Build a minimal single-page PDF with extractable text (no binary fixture)."""
    # Escape parentheses and backslashes for PDF literal strings.
    safe = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 50 50 Td ({safe}) Tj ET"
    stream_bytes = stream.encode("latin-1", errors="replace")
    objects = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        (
            b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 400 200] "
            b"/Contents 4 0 R /Resources<< /Font<< /F1 5 0 R >> >> >>endobj\n"
        ),
        (
            f"4 0 obj<< /Length {len(stream_bytes)} >>stream\n".encode("ascii")
            + stream_bytes
            + b"\nendstream\nendobj\n"
        ),
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
    ]
    header = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
    body = b"".join(objects)
    offsets = [0]
    cursor = len(header)
    for obj in objects:
        offsets.append(cursor)
        cursor += len(obj)
    xref_offset = cursor
    xref_lines = [b"xref\n0 6\n", b"0000000000 65535 f \n"]
    for offset in offsets[1:]:
        xref_lines.append(f"{offset:010d} 00000 n \n".encode("ascii"))
    trailer = (
        b"trailer<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n" + f"{xref_offset}\n".encode("ascii") + b"%%EOF\n"
    )
    return header + body + b"".join(xref_lines) + trailer


def make_empty_pdf_bytes() -> bytes:
    """PDF with a blank page and no extractable text."""
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_invalid_pdf_bytes() -> bytes:
    return b"%PDF-1.4\n% corrupted truncated pdf"


def make_docx_bytes(text: str = "Cortexa is a local-first AI agent platform.") -> bytes:
    document = DocxDocument()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_empty_docx_bytes() -> bytes:
    document = DocxDocument()
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
